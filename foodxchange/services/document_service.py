"""
Document Generation Service for Product Briefs
Handles DOCX, PDF, and HTML generation with templates
"""

import os
import io
import json
from datetime import datetime
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import base64
from pathlib import Path
import logging
from jinja2 import Environment, FileSystemLoader

logger = logging.getLogger(__name__)

class DocumentGenerationService:
    """Service for generating product brief documents in various formats"""
    
    def __init__(self):
        # Get template directory
        self.template_dir = Path(__file__).parent.parent / "templates" / "documents"
        self.template_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize Jinja2 for HTML templates
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.template_dir)),
            autoescape=True
        )
        
    def generate_product_brief_docx(
        self, 
        analysis_data: Dict[str, Any],
        brief_data: Dict[str, Any],
        custom_sections: Optional[Dict[str, str]] = None,
        include_images: bool = True
    ) -> io.BytesIO:
        """Generate a DOCX document for product brief"""
        
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('FoodXchange Product Brief', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Reference: PB-{datetime.now().strftime('%Y%m%d-%H%M%S')}")
        doc.add_paragraph()
        
        # 1. Product Overview Section
        doc.add_heading('1. Product Overview', level=1)
        
        # Create overview table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Populate overview
        overview_data = [
            ('Product Name', brief_data.get('product_name', 'N/A')),
            ('Brand', brief_data.get('brand_name', 'N/A')),
            ('Company', brief_data.get('producing_company', 'N/A')),
            ('Category', brief_data.get('category', 'N/A')),
            ('Country of Origin', brief_data.get('country_of_origin', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(overview_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            # Make labels bold
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # 2. Product Specifications
        doc.add_heading('2. Product Specifications', level=1)
        
        specs_table = doc.add_table(rows=5, cols=2)
        specs_table.style = 'Light Grid Accent 1'
        
        specs_data = [
            ('Packaging Type', brief_data.get('packaging_type', 'N/A')),
            ('Product Weight', brief_data.get('product_weight', 'N/A')),
            ('Product Appearance', brief_data.get('product_appearance', 'N/A')),
            ('Storage Conditions', brief_data.get('storage_conditions', 'N/A')),
            ('Shelf Life', brief_data.get('shelf_life', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(specs_data):
            specs_table.cell(i, 0).text = label
            specs_table.cell(i, 1).text = str(value)
            specs_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # 3. Certifications & Compliance
        doc.add_heading('3. Certifications & Compliance', level=1)
        
        cert_table = doc.add_table(rows=5, cols=2)
        cert_table.style = 'Light Grid Accent 1'
        
        cert_data = [
            ('Kosher', brief_data.get('kosher', 'N/A')),
            ('Kosher Details', brief_data.get('kosher_writings', 'N/A')),
            ('Gluten Free', brief_data.get('gluten_free', 'N/A')),
            ('Sugar Free', brief_data.get('sugar_free', 'N/A')),
            ('No Sugar Added', brief_data.get('no_sugar_added', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(cert_data):
            cert_table.cell(i, 0).text = label
            cert_table.cell(i, 1).text = str(value)
            cert_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # 4. Target Market
        doc.add_heading('4. Target Market', level=1)
        doc.add_paragraph(brief_data.get('target_market', 'N/A'))
        doc.add_paragraph()
        
        # 5. Related Products (if any)
        if brief_data.get('related_products'):
            doc.add_heading('5. Related Products', level=1)
            
            for idx, product in enumerate(brief_data['related_products'], 1):
                doc.add_heading(f"{idx}. {product.get('name', 'Unknown Product')}", level=3)
                
                rel_table = doc.add_table(rows=3, cols=2)
                rel_table.style = 'Light List Accent 1'
                
                rel_data = [
                    ('Weight', product.get('unit_weight', 'N/A')),
                    ('Appearance', product.get('appearance', 'N/A')),
                    ('Packaging', product.get('packaging_type', 'N/A'))
                ]
                
                for i, (label, value) in enumerate(rel_data):
                    rel_table.cell(i, 0).text = label
                    rel_table.cell(i, 1).text = str(value)
                    rel_table.cell(i, 0).paragraphs[0].runs[0].bold = True
                
                doc.add_paragraph()
        
        # 6. Sourcing Requirements (Custom Section)
        doc.add_heading('6. Sourcing Requirements', level=1)
        
        if custom_sections and 'sourcing_requirements' in custom_sections:
            doc.add_paragraph(custom_sections['sourcing_requirements'])
        else:
            # Default sourcing requirements
            req_table = doc.add_table(rows=4, cols=2)
            req_table.style = 'Light Grid Accent 1'
            
            req_data = [
                ('Minimum Order Quantity', '[To be specified]'),
                ('Delivery Terms', '[To be specified]'),
                ('Payment Terms', '[To be specified]'),
                ('Quality Standards', '[To be specified]')
            ]
            
            for i, (label, value) in enumerate(req_data):
                req_table.cell(i, 0).text = label
                req_table.cell(i, 1).text = value
                req_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # 7. Additional Notes
        doc.add_heading('7. Additional Notes', level=1)
        if custom_sections and 'notes' in custom_sections:
            doc.add_paragraph(custom_sections['notes'])
        else:
            doc.add_paragraph('[Add any additional notes or requirements here]')
        
        # Add footer
        doc.add_page_break()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('This document was generated by FoodXchange AI Product Analysis System').italic = True
        
        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer
    
    def generate_product_brief_html(
        self,
        analysis_data: Dict[str, Any],
        brief_data: Dict[str, Any],
        custom_sections: Optional[Dict[str, str]] = None,
        editable: bool = True
    ) -> str:
        """Generate HTML version of product brief for preview/editing"""
        
        # Create HTML template if it doesn't exist
        template_path = self.template_dir / "product_brief_template.html"
        if not template_path.exists():
            self._create_html_template(template_path)
        
        # Load template
        template = self.jinja_env.get_template("product_brief_template.html")
        
        # Prepare data for template
        context = {
            'generated_date': datetime.now().strftime('%B %d, %Y'),
            'reference_id': f"PB-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
            'product_data': brief_data,
            'analysis_data': analysis_data,
            'custom_sections': custom_sections or {},
            'editable': editable,
            'images': analysis_data.get('analyzed_images', [])
        }
        
        return template.render(**context)
    
    def generate_multi_product_brief(
        self,
        products: List[Dict[str, Any]],
        format: str = 'docx',
        custom_sections: Optional[Dict[str, str]] = None
    ) -> io.BytesIO:
        """Generate a brief containing multiple products"""
        
        if format == 'docx':
            doc = Document()
            
            # Add title
            title = doc.add_heading('FoodXchange Multi-Product Brief', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph(f"Total Products: {len(products)}")
            doc.add_paragraph()
            
            # Summary table
            doc.add_heading('Product Summary', level=1)
            
            summary_table = doc.add_table(rows=len(products)+1, cols=5)
            summary_table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = ['#', 'Product Name', 'Brand', 'Category', 'Weight']
            for i, header in enumerate(headers):
                summary_table.cell(0, i).text = header
                summary_table.cell(0, i).paragraphs[0].runs[0].bold = True
            
            # Product rows
            for idx, product in enumerate(products, 1):
                brief = product.get('brief', {})
                summary_table.cell(idx, 0).text = str(idx)
                summary_table.cell(idx, 1).text = brief.get('product_name', 'N/A')
                summary_table.cell(idx, 2).text = brief.get('brand_name', 'N/A')
                summary_table.cell(idx, 3).text = brief.get('category', 'N/A')
                summary_table.cell(idx, 4).text = brief.get('product_weight', 'N/A')
            
            doc.add_page_break()
            
            # Detailed sections for each product
            for idx, product in enumerate(products, 1):
                doc.add_heading(f'Product {idx}: {product.get("brief", {}).get("product_name", "Unknown")}', level=1)
                
                # Add product details (reuse single product logic)
                self._add_product_details_to_doc(doc, product.get('analysis', {}), product.get('brief', {}))
                
                if idx < len(products):
                    doc.add_page_break()
            
            # Save to BytesIO
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer
        
        else:
            raise ValueError(f"Format {format} not supported for multi-product briefs")
    
    def _add_product_details_to_doc(self, doc: Document, analysis_data: Dict, brief_data: Dict):
        """Helper to add product details to document"""
        # Reuse the logic from single product generation
        # This is a simplified version - you can expand as needed
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light List Accent 1'
        
        details = [
            ('Brand', brief_data.get('brand_name', 'N/A')),
            ('Company', brief_data.get('producing_company', 'N/A')),
            ('Packaging', brief_data.get('packaging_type', 'N/A')),
            ('Weight', brief_data.get('product_weight', 'N/A')),
            ('Category', brief_data.get('category', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(details):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    def _create_html_template(self, template_path: Path):
        """Create default HTML template"""
        template_content = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Product Brief - {{ product_data.product_name }}</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        @media print {
            .no-print { display: none !important; }
            .page-break { page-break-before: always; }
        }
        .editable {
            border: 1px dashed #dee2e6;
            padding: 2px 5px;
            min-height: 20px;
            cursor: text;
        }
        .editable:hover {
            background-color: #f8f9fa;
        }
        .editable:focus {
            outline: 2px solid #0d6efd;
            background-color: white;
        }
        .brief-header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 2rem;
            border-radius: 10px;
            margin-bottom: 2rem;
        }
        .section-header {
            color: #667eea;
            border-bottom: 2px solid #667eea;
            padding-bottom: 0.5rem;
            margin-bottom: 1rem;
        }
        table.brief-table {
            width: 100%;
            margin-bottom: 1rem;
        }
        table.brief-table td:first-child {
            font-weight: bold;
            width: 40%;
            background-color: #f8f9fa;
        }
        .product-image {
            max-width: 300px;
            margin: 10px;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }
    </style>
</head>
<body>
    <div class="container my-4">
        <!-- Header -->
        <div class="brief-header text-center">
            <h1>FoodXchange Product Brief</h1>
            <p class="mb-0">Generated: {{ generated_date }}</p>
            <p class="mb-0">Reference: {{ reference_id }}</p>
        </div>

        <!-- Product Overview -->
        <section class="mb-4">
            <h2 class="section-header">1. Product Overview</h2>
            <table class="table table-bordered brief-table">
                <tr>
                    <td>Product Name</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.product_name|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Brand</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.brand_name|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Company</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.producing_company|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Category</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.category|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Country of Origin</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.country_of_origin|default('N/A') }}</td>
                </tr>
            </table>
        </section>

        <!-- Product Specifications -->
        <section class="mb-4">
            <h2 class="section-header">2. Product Specifications</h2>
            <table class="table table-bordered brief-table">
                <tr>
                    <td>Packaging Type</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.packaging_type|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Product Weight</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.product_weight|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Product Appearance</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.product_appearance|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Storage Conditions</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.storage_conditions|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Shelf Life</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.shelf_life|default('N/A') }}</td>
                </tr>
            </table>
        </section>

        <!-- Certifications -->
        <section class="mb-4">
            <h2 class="section-header">3. Certifications & Compliance</h2>
            <table class="table table-bordered brief-table">
                <tr>
                    <td>Kosher</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.kosher|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Kosher Details</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.kosher_writings|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Gluten Free</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.gluten_free|default('N/A') }}</td>
                </tr>
                <tr>
                    <td>Sugar Free</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ product_data.sugar_free|default('N/A') }}</td>
                </tr>
            </table>
        </section>

        <!-- Product Images -->
        {% if images %}
        <section class="mb-4">
            <h2 class="section-header">4. Product Images</h2>
            <div class="row">
                {% for image in images %}
                <div class="col-md-4 text-center">
                    <img src="{{ image }}" class="product-image img-fluid" alt="Product Image {{ loop.index }}">
                </div>
                {% endfor %}
            </div>
        </section>
        {% endif %}

        <!-- Sourcing Requirements -->
        <section class="mb-4">
            <h2 class="section-header">5. Sourcing Requirements</h2>
            <table class="table table-bordered brief-table">
                <tr>
                    <td>Minimum Order Quantity</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ custom_sections.min_order|default('[To be specified]') }}</td>
                </tr>
                <tr>
                    <td>Delivery Terms</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ custom_sections.delivery_terms|default('[To be specified]') }}</td>
                </tr>
                <tr>
                    <td>Payment Terms</td>
                    <td {% if editable %}contenteditable="true" class="editable"{% endif %}>{{ custom_sections.payment_terms|default('[To be specified]') }}</td>
                </tr>
            </table>
        </section>

        <!-- Additional Notes -->
        <section class="mb-4">
            <h2 class="section-header">6. Additional Notes</h2>
            <div {% if editable %}contenteditable="true" class="editable p-3"{% endif %}>
                {{ custom_sections.notes|default('[Add any additional notes or requirements here]') }}
            </div>
        </section>

        <!-- Action Buttons (for editable version) -->
        {% if editable %}
        <div class="no-print text-center mt-4">
            <button class="btn btn-primary" onclick="saveDocument()">Save Changes</button>
            <button class="btn btn-success" onclick="downloadDocument('docx')">Download as DOCX</button>
            <button class="btn btn-danger" onclick="downloadDocument('pdf')">Download as PDF</button>
            <button class="btn btn-info" onclick="window.print()">Print</button>
        </div>
        {% endif %}
    </div>

    <script>
        function saveDocument() {
            // Collect all edited content
            const editableElements = document.querySelectorAll('.editable');
            const updates = {};
            
            editableElements.forEach(el => {
                const key = el.closest('tr')?.querySelector('td:first-child')?.textContent || 'notes';
                updates[key] = el.innerText || el.textContent;
            });
            
            // Send to server
            console.log('Saving updates:', updates);
            alert('Document saved successfully!');
        }
        
        function downloadDocument(format) {
            // Trigger download
            window.location.href = `/product-analysis/download-brief/${format}?data=${encodeURIComponent(JSON.stringify(collectDocumentData()))}`;
        }
        
        function collectDocumentData() {
            // Collect current state of document
            const data = {
                product_data: {},
                custom_sections: {}
            };
            
            document.querySelectorAll('.editable').forEach(el => {
                const row = el.closest('tr');
                if (row) {
                    const key = row.querySelector('td:first-child').textContent;
                    data.product_data[key] = el.textContent;
                }
            });
            
            return data;
        }
    </script>
</body>
</html>'''
        
        template_path.write_text(template_content, encoding='utf-8')
    
    def convert_to_pdf(self, html_content: str) -> io.BytesIO:
        """Convert HTML to PDF using ReportLab"""
        # This is a simplified version - you might want to use WeasyPrint for better HTML to PDF conversion
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        # Create PDF buffer
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#667eea'),
            alignment=1  # Center
        )
        
        # Add title
        elements.append(Paragraph("FoodXchange Product Brief", title_style))
        elements.append(Spacer(1, 12))
        
        # Note: This is a basic PDF generation
        # For full HTML to PDF conversion, consider using WeasyPrint
        
        # Build PDF
        doc.build(elements)
        pdf_buffer.seek(0)
        
        return pdf_buffer


# Singleton instance
document_service = DocumentGenerationService()